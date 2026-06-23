#!/usr/bin/env python3
import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 1. Load posts ──────────────────────────────────────────────────────────────
with open('/data/data/com.termux/files/home/your_posts.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

posts = []
for post in data:
    if 'data' in post:
        for d in post['data']:
            if 'post' in d:
                text = d['post'].encode('latin-1').decode('utf-8')
                if len(text) >= 300:
                    posts.append({'ts': post.get('timestamp', 0), 'text': text})
                break

posts.sort(key=lambda x: x['ts'])
print(f"Loaded {len(posts)} posts")

# ── 2. Classify each post ──────────────────────────────────────────────────────
CHAPTERS = [
    {
        'name': 'על הכתיבה',
        'keywords': ['כתיבה','כותב','לכתוב','מחברת','פוסט','כתב את','הכתיבה עושה','לכתוב כל יום',
                     'לכתוב בחוץ','לכתוב בחוף','הכתיבה שלי','כתיבה משחררת','כתיבה גורמת',
                     'לכתוב ברגל','פסיכולוג הציע לכתוב'],
    },
    {
        'name': 'הרגישות שלי',
        'keywords': ['רגיש מאוד','אדם רגיש','רגישות','HSP','הספר אדם רגיש','הרגשות שלי',
                     'מרגיש מותקף','מרגיש פגוע','רגש','שפת הרגש','כאב לי','הכאיב לי',
                     'קשה לי להרגיש','הרגשה קשה','מרגיש רע מעצמי'],
    },
    {
        'name': 'עבודה עם ילדים מיוחדים',
        'keywords': ['אוטיסט','ילדים מיוחדים','דיירים','דייר','מדריך','בית הספר','הוסטל',
                     'רכז','הילד','ילד אחד','ילד חדש','בית ספר','דירה ב','דירה ק',
                     'ילד עם חרדות','הילדים','ישנים בדירה'],
    },
    {
        'name': 'משפחה',
        'keywords': ['אימא שלי','אבא שלי','אח שלי','אחות שלי','הורים שלי','משפחה שלי',
                     'אחים שלי','אחותי','בית ההורים','ליל שבת','סעודה','משפחה',
                     'אחי','האחים','בבית ההורים'],
    },
    {
        'name': 'בדידות וחיפוש קשר',
        'keywords': ['בדידות','בודד','חסר','חסך','חיבוק','אין לי חברים','אין לי עם מי',
                     'מתגעגע','להכיר אנשים','קשר זוגי','בחורה','אהבה','בודד בעולם',
                     'הרגשת חוסר','בודדות'],
    },
    {
        'name': 'ימים קשים',
        'keywords': ['למות','לא רוצה לקום','מחשבות שליליות','דיכאון','יאוש','לא רוצה לחיות',
                     'עדיף לא להיות קיים','רוצה למות','חשבתי למות','חשבתי שעדיף',
                     'אין לי סיבה לחיות','אין למה לקום','אין סיבה לחיים','ימים קשים',
                     'מחשבות קשות','חושב שאין לי תיקוה'],
    },
    {
        'name': 'אמונה ורוחניות',
        'keywords': ['ברסלב','ליקוטי','ישיבה','שבת','תורה','תפילה','בעזרת השם','ברוך השם',
                     'הרמב"ם','יהדות','אמונה','השם','ספר מוסר','חגים','ראש השנה',
                     'יום כיפור','סוכות','חנוכה','ליקוטי עיצות','ליקוטי עצות'],
    },
    {
        'name': 'עצמאות ועתיד',
        'keywords': ['לעזוב את הבית','עצמאות','ללמוד אנגלית','סייבר','האקינג','לימודים',
                     'עתיד שלי','תוכניות','לגדול','חלומות','רוצה להתקדם','לקבל דירה',
                     'לצאת מבית ההורים','מגורים','עסק'],
    },
    {
        'name': 'מחשבות על החיים',
        'keywords': ['מחשבה על','חשבתי על','להתבונן','חשבון נפש','הבנתי ש','מסקנה',
                     'הרהורים','תחושה ש','עניין של','לחשוב על זה','חשבתי לעצמי',
                     'מה שלמדתי','נרקסיזם','פסיכולוגיה','דפוסי חשיבה','מחשבות על'],
    },
]
DEFAULT_CHAPTER = 'יומיום'

def score(text, kws):
    t = text
    return sum(t.count(k) for k in kws)

# assign each post to best-matching chapter
for p in posts:
    scores = [score(p['text'], ch['keywords']) for ch in CHAPTERS]
    best = max(range(len(scores)), key=lambda i: scores[i])
    p['chapter'] = CHAPTERS[best]['name'] if scores[best] > 0 else DEFAULT_CHAPTER

# print stats
from collections import Counter
counts = Counter(p['chapter'] for p in posts)
for ch in CHAPTERS:
    print(f"  {ch['name']}: {counts.get(ch['name'],0)}")
print(f"  {DEFAULT_CHAPTER}: {counts.get(DEFAULT_CHAPTER,0)}")

# ── 3. Build Word document ─────────────────────────────────────────────────────
doc = Document()

# Page setup: RTL, A4
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin  = Cm(3)
section.right_margin = Cm(3)
section.top_margin   = Cm(2.5)
section.bottom_margin = Cm(2.5)

def set_rtl_para(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)

def add_chapter_title(doc, name):
    para = doc.add_paragraph()
    set_rtl_para(para)
    run = para.add_run(name)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x8c)
    para.paragraph_format.space_before = Pt(24)
    para.paragraph_format.space_after  = Pt(12)

def add_post(doc, text):
    # separator line
    sep = doc.add_paragraph()
    set_rtl_para(sep)
    sep.add_run('─' * 40)
    sep.paragraph_format.space_before = Pt(10)
    sep.paragraph_format.space_after  = Pt(4)

    # post body
    para = doc.add_paragraph()
    set_rtl_para(para)
    run = para.add_run(text)
    run.font.size = Pt(12)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(8)

# Title page
title_para = doc.add_paragraph()
set_rtl_para(title_para)
title_run = title_para.add_run('ספר בצלאל')
title_run.bold = True
title_run.font.size = Pt(32)
title_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x8c)
title_para.paragraph_format.space_before = Pt(100)
title_para.paragraph_format.space_after  = Pt(20)
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Chapters
chapter_order = [ch['name'] for ch in CHAPTERS] + [DEFAULT_CHAPTER]
for ch_name in chapter_order:
    chapter_posts = [p for p in posts if p['chapter'] == ch_name]
    if not chapter_posts:
        continue
    add_chapter_title(doc, ch_name)
    for p in chapter_posts:
        add_post(doc, p['text'])
    doc.add_page_break()

out_path = '/data/data/com.termux/files/home/sefer_bezalel.docx'
doc.save(out_path)
print(f"\nSaved: {out_path}")
print(f"Total posts in book: {len(posts)}")
